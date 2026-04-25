from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_presentation_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('MissionControl: Agentic AI for Autonomous Enterprise Workflows', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Automating tasks, orchestrating agents, and streamlining operations.')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Slide 1
    doc.add_heading('Slide 1: Title Slide', level=1)
    doc.add_paragraph('Title: MissionControl: Agentic AI for Autonomous Enterprise Workflows')
    doc.add_paragraph('Subtitle: Automating tasks, orchestrating agents, and streamlining operations.')
    doc.add_paragraph('Key Elements to Include:')
    doc.add_paragraph('• Project Name/Logo', style='List Bullet')
    doc.add_paragraph('• Date & Presenter Name', style='List Bullet')
    doc.add_paragraph('• A bold, clean aesthetic matching the application\'s design', style='List Bullet')
    doc.add_page_break()

    # Slide 2
    doc.add_heading('Slide 2: The Vision & Problem Solved', level=1)
    doc.add_paragraph('Title: Transforming Enterprise Automation')
    doc.add_paragraph('Content:')
    doc.add_paragraph('• The Problem: Manual task delegation, fragmented communication, and lack of oversight in complex workflows lead to bottlenecks.', style='List Bullet')
    doc.add_paragraph('• The Solution: MissionControl acts as an intelligent operating system that interprets natural language (transcripts) and autonomously executes multi-step workflows.', style='List Bullet')
    doc.add_paragraph('• Core Value: Reduces cognitive load, automates cross-platform actions, and ensures human-in-the-loop safety for high-risk actions.', style='List Bullet')
    doc.add_page_break()

    # Slide 3
    doc.add_heading('Slide 3: System Architecture & Tech Stack', level=1)
    doc.add_paragraph('Title: Robust & Scalable Foundation')
    doc.add_paragraph('Content:')
    doc.add_paragraph('• Frontend: React + Vite ecosystem for a dynamic, real-time user interface.', style='List Bullet')
    doc.add_paragraph('• Backend: Python + FastAPI for high-performance API routing and background task processing.', style='List Bullet')
    doc.add_paragraph('• AI Orchestration: LangGraph powers the stateful, multi-agent pipeline.', style='List Bullet')
    doc.add_paragraph('• LLM Engine: Multi-provider LLM support with built-in fallback mechanisms for maximum reliability and uptime.', style='List Bullet')
    doc.add_page_break()

    # Slide 4
    doc.add_heading('Slide 4: The Agentic Workflow Pipeline', level=1)
    doc.add_paragraph('Title: Intelligent Multi-Agent Orchestration')
    doc.add_paragraph('Content:')
    doc.add_paragraph('Our pipeline processes unstructured transcripts through specialized AI agents:')
    doc.add_paragraph('1. Scribe: Transcribes and parses the initial input into actionable context.', style='List Number')
    doc.add_paragraph('2. Planner: Breaks down the context into structured, logical tasks.', style='List Number')
    doc.add_paragraph('3. Executor: Carries out the tasks (sends emails, schedules events).', style='List Number')
    doc.add_paragraph('4. Auditor: Verifies task completion and checks for errors or quality issues.', style='List Number')
    doc.add_paragraph('5. Escalation Check: Flags high-risk operations for human review before proceeding.', style='List Number')
    doc.add_page_break()

    # Slide 5
    doc.add_heading('Slide 5: Standout "Story" Features (UI/UX)', level=1)
    doc.add_paragraph('Title: Selling the Vision: The Demo Moments')
    doc.add_paragraph('Content:')
    doc.add_paragraph('• Workflow Risk Score (Pre-Launch): Analyzes transcript constraints before execution. Example: "Risk: 73/100 — Ravi is overloaded, Sept 21 has 2 clashing deadlines." Sets up the problem for the AI to solve.', style='List Bullet')
    doc.add_paragraph('• Cognitive Load Monitor (Real Product Moment): Live visualizations of team bandwidth. Automatically reassigns tasks (e.g., from Ravi at 90% load to James) with clear AI reasoning.', style='List Bullet')
    doc.add_paragraph('• Workflow Obituary (Polish Moment): Elegant, formal permanent records for permanently failed tasks. Reads like an executive report: "The primary strategy failed... Recommended action: reassign to a senior resource."', style='List Bullet')
    doc.add_page_break()

    # Slide 6
    doc.add_heading('Slide 6: Seamless External Integrations', level=1)
    doc.add_paragraph('Title: Connecting with the Real World')
    doc.add_paragraph('Content:')
    doc.add_paragraph('• Google Workspace OAuth2 Integration: Secure, automated authentication flow.', style='List Bullet')
    doc.add_paragraph('• Automated Gmail Actions: Dynamically drafts and sends task assignment and notification emails to team members.', style='List Bullet')
    doc.add_paragraph('• Google Calendar Sync: Automatically schedules events, meetings, and deadlines parsed from the workflow transcript.', style='List Bullet')
    doc.add_paragraph('• Real-time Status Panels: Dedicated UI components tracking the live status of Emails and Calendar integrations.', style='List Bullet')
    doc.add_page_break()

    # Slide 7
    doc.add_heading('Slide 7: Summary & Future Outlook', level=1)
    doc.add_paragraph('Title: The Future of MissionControl')
    doc.add_paragraph('Content:')
    doc.add_paragraph('• Summary: Successfully built an end-to-end, autonomous workflow OS that understands context, acts on it, and handles errors gracefully.', style='List Bullet')
    doc.add_paragraph('• Stability Achieved: Premium "nude" color palette dashboard, robust LLM connectivity fallbacks, and multi-agent pipeline.', style='List Bullet')
    doc.add_paragraph('• Next Steps: Expanding the integration ecosystem (e.g., Slack, Jira) and further refining the autonomous decision-making confidence.', style='List Bullet')
    doc.add_paragraph('• Call to Action: Q&A / Live Demo.', style='List Bullet')

    doc.save('c:/et/mission-control/MissionControl_Presentation.docx')
    print("Document saved successfully at c:/et/mission-control/MissionControl_Presentation.docx")

if __name__ == "__main__":
    create_presentation_doc()
